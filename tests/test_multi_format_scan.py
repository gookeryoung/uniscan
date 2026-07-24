"""多格式扫描与详情定位集成测试。

覆盖典型文件格式（txt/yaml/json/docx/xlsx/odt/zip/二进制）中
数据库连接串与 Bearer 令牌的扫描与 ``match_text`` 传递验证。

测试策略：
- 动态生成各格式测试文件，避免二进制 fixture 入仓
- 使用内置规则集端到端验证：文件 → 提取器 → 扫描器 → RuleHit.match_text
- 不支持生成的格式（pdf/wps/rar）由对应单元测试覆盖，此处不重复
"""

from __future__ import annotations

import os
import zipfile
from pathlib import Path
from typing import Sequence

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

from fuscan.config import load_builtin_ruleset
from fuscan.rules.model import Rule, RuleSet, Severity
from fuscan.scanner import Scanner
from fuscan.scanner.result import RuleHit, ScanResult

# 数据库连接串与 Bearer 令牌的测试样本
_DB_CONN_SAMPLE = "mongodb://user:pass123@host:27017/db"
_BEARER_SAMPLE = "Bearer eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiIxMjMifQ.sig"
_DB_CONN_WITH_BACKSLASH = r"mongodb://user:pass\123@host"
_DB_CONN_WITH_QUOTE = "mongodb://user:pa'ss@host"
_BEARER_CROSS_LINE = "Bearer\n  eyJhbGci.token"

# 内置规则集端到端测试样本（对应 assets/rules/builtin.yaml 中的 P0101/P0102 规则）
_PASSWORD_SAMPLE = "password=S3cr3t!"
_PRIVATE_KEY_SAMPLE = "-----BEGIN RSA PRIVATE KEY-----"


def _db_rule() -> Rule:
    """构造数据库连接串规则（与内置规则一致，独立使用避免依赖内置规则加载顺序）。"""
    from fuscan.rules.model import LeafMatch, MatchMode, MatchTarget

    return Rule(
        name="数据库连接串",
        severity=Severity.WARNING,
        match=LeafMatch(
            target=MatchTarget.CONTENT,
            mode=MatchMode.REGEX,
            pattern=r"(?i)(mongodb|postgres|postgresql|mysql|redis)://\S+:\S+@",
        ),
    )


def _bearer_rule() -> Rule:
    """构造 Bearer 令牌规则。"""
    from fuscan.rules.model import LeafMatch, MatchMode, MatchTarget

    return Rule(
        name="Bearer令牌",
        severity=Severity.INFO,
        match=LeafMatch(
            target=MatchTarget.CONTENT,
            mode=MatchMode.REGEX,
            pattern=r"(?i)bearer\s+[A-Za-z0-9._\-]+",
        ),
    )


def _build_ruleset(*rules: Rule) -> RuleSet:
    return RuleSet(version="1.0", rules=tuple(rules))


def _find_hit(report_results: Sequence[ScanResult], rule_name: str) -> RuleHit | None:
    """从扫描结果中查找指定规则的命中。"""
    for sr in report_results:
        for hit in sr.hits:
            if hit.rule_name == rule_name:
                return hit
    return None


# ---------------------------------------------------------------------------
# 文本格式：txt / yaml / json
# ---------------------------------------------------------------------------


class TestTextFormats:
    """txt/yaml/json 文本格式扫描与 match_text 验证。"""

    def test_txt_db_connection_match_text(self, tmp_path: Path) -> None:
        """txt 文件中的数据库连接串应被扫描到且 match_text 为原始匹配文本。"""
        path = tmp_path / "config.txt"
        path.write_text(f"db_url = {_DB_CONN_SAMPLE}\n", encoding="utf-8")
        scanner = Scanner(_build_ruleset(_db_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert hit.match_text == "mongodb://user:pass123@"

    def test_txt_bearer_match_text(self, tmp_path: Path) -> None:
        """txt 文件中的 Bearer 令牌应被扫描到且 match_text 为原始匹配文本。"""
        path = tmp_path / "auth.txt"
        path.write_text(f"Authorization: {_BEARER_SAMPLE}\n", encoding="utf-8")
        scanner = Scanner(_build_ruleset(_bearer_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "Bearer令牌")
        assert hit is not None
        assert hit.match_text == _BEARER_SAMPLE

    def test_yaml_db_connection_match_text(self, tmp_path: Path) -> None:
        """yaml 文件中的数据库连接串应被扫描到且 match_text 正确。"""
        path = tmp_path / "app.yaml"
        path.write_text(
            f"database:\n  url: {_DB_CONN_SAMPLE}\n  pool: 10\n",
            encoding="utf-8",
        )
        scanner = Scanner(_build_ruleset(_db_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert "mongodb://user:pass123@" in hit.match_text

    def test_json_bearer_match_text(self, tmp_path: Path) -> None:
        """json 文件中的 Bearer 令牌应被扫描到且 match_text 正确。"""
        path = tmp_path / "tokens.json"
        path.write_text(
            f'{{"auth": "Authorization: {_BEARER_SAMPLE}", "timeout": 30}}',
            encoding="utf-8",
        )
        scanner = Scanner(_build_ruleset(_bearer_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "Bearer令牌")
        assert hit is not None
        assert hit.match_text == _BEARER_SAMPLE

    def test_txt_db_connection_with_backslash(self, tmp_path: Path) -> None:
        """密码含反斜杠的数据库连接串 match_text 应原样保留反斜杠。"""
        path = tmp_path / "db.txt"
        path.write_text(f"url = {_DB_CONN_WITH_BACKSLASH}\n", encoding="utf-8")
        scanner = Scanner(_build_ruleset(_db_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert "\\" in hit.match_text
        assert hit.match_text == r"mongodb://user:pass\123@"

    def test_txt_db_connection_with_single_quote(self, tmp_path: Path) -> None:
        """密码含单引号的数据库连接串 match_text 应原样保留单引号。"""
        path = tmp_path / "db.txt"
        path.write_text(f"url = {_DB_CONN_WITH_QUOTE}\n", encoding="utf-8")
        scanner = Scanner(_build_ruleset(_db_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert "'" in hit.match_text

    def test_txt_bearer_cross_line(self, tmp_path: Path) -> None:
        """跨行 Bearer 令牌的 match_text 应保留换行符。"""
        path = tmp_path / "auth.txt"
        path.write_text(f"Authorization: {_BEARER_CROSS_LINE}\n", encoding="utf-8")
        scanner = Scanner(_build_ruleset(_bearer_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "Bearer令牌")
        assert hit is not None
        assert "\n" in hit.match_text

    def test_txt_no_false_positive_on_normal_content(self, tmp_path: Path) -> None:
        """普通文本不应误报数据库连接串或 Bearer 令牌。"""
        path = tmp_path / "normal.txt"
        path.write_text("这是一段普通文本，不包含任何敏感信息。\nhello world\n", encoding="utf-8")
        scanner = Scanner(_build_ruleset(_db_rule(), _bearer_rule()))
        report = scanner.scan(tmp_path)
        assert len(report.hits) == 0


# ---------------------------------------------------------------------------
# 文档格式：docx / xlsx / odt
# ---------------------------------------------------------------------------


class TestDocumentFormats:
    """docx/xlsx/odt 文档格式扫描与 match_text 验证。"""

    def test_docx_db_connection_match_text(self, tmp_path: Path) -> None:
        """docx 文档中的数据库连接串应被扫描到且 match_text 正确。"""
        from docx import Document

        doc = Document()
        doc.add_paragraph(f"数据库配置: {_DB_CONN_SAMPLE}")
        doc.add_paragraph("普通段落")
        path = tmp_path / "config.docx"
        doc.save(str(path))

        scanner = Scanner(_build_ruleset(_db_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert "mongodb://user:pass123@" in hit.match_text

    def test_docx_bearer_match_text(self, tmp_path: Path) -> None:
        """docx 文档中的 Bearer 令牌应被扫描到且 match_text 正确。"""
        from docx import Document

        doc = Document()
        doc.add_paragraph(f"认证头: Authorization: {_BEARER_SAMPLE}")
        path = tmp_path / "auth.docx"
        doc.save(str(path))

        scanner = Scanner(_build_ruleset(_bearer_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "Bearer令牌")
        assert hit is not None
        assert hit.match_text == _BEARER_SAMPLE

    def test_xlsx_db_connection_match_text(self, tmp_path: Path) -> None:
        """xlsx 表格中的数据库连接串应被扫描到且 match_text 正确。"""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        assert ws is not None
        ws["A1"] = "连接串"
        ws["B1"] = _DB_CONN_SAMPLE
        path = tmp_path / "data.xlsx"
        wb.save(str(path))

        scanner = Scanner(_build_ruleset(_db_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert "mongodb://user:pass123@" in hit.match_text

    def test_xlsx_bearer_match_text(self, tmp_path: Path) -> None:
        """xlsx 表格中的 Bearer 令牌应被扫描到且 match_text 正确。"""
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        assert ws is not None
        ws["A1"] = _BEARER_SAMPLE
        path = tmp_path / "tokens.xlsx"
        wb.save(str(path))

        scanner = Scanner(_build_ruleset(_bearer_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "Bearer令牌")
        assert hit is not None
        assert hit.match_text == _BEARER_SAMPLE

    def test_odt_db_connection_match_text(self, tmp_path: Path) -> None:
        """odt 文档中的数据库连接串应被扫描到且 match_text 正确。"""
        from odf.opendocument import OpenDocumentText
        from odf.text import P

        doc = OpenDocumentText()
        p = P(text=f"数据库连接: {_DB_CONN_SAMPLE}")
        doc.text.addElement(p)  # pyrefly: ignore [missing-attribute]
        path = tmp_path / "config.odt"
        doc.save(str(path))

        scanner = Scanner(_build_ruleset(_db_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert "mongodb://user:pass123@" in hit.match_text

    def test_odt_bearer_match_text(self, tmp_path: Path) -> None:
        """odt 文档中的 Bearer 令牌应被扫描到且 match_text 正确。"""
        from odf.opendocument import OpenDocumentText
        from odf.text import P

        doc = OpenDocumentText()
        p = P(text=f"Authorization: {_BEARER_SAMPLE}")
        doc.text.addElement(p)  # pyrefly: ignore [missing-attribute]
        path = tmp_path / "auth.odt"
        doc.save(str(path))

        scanner = Scanner(_build_ruleset(_bearer_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "Bearer令牌")
        assert hit is not None
        assert hit.match_text == _BEARER_SAMPLE


# ---------------------------------------------------------------------------
# 压缩格式：zip
# ---------------------------------------------------------------------------


class TestArchiveFormats:
    """zip 压缩格式扫描与 match_text 验证。"""

    def test_zip_txt_db_connection_match_text(self, tmp_path: Path) -> None:
        """zip 内 txt 文件的数据库连接串应被扫描到且 match_text 正确。"""
        zip_path = tmp_path / "archive.zip"
        with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("config.txt", f"db_url = {_DB_CONN_SAMPLE}\n")

        scanner = Scanner(_build_ruleset(_db_rule()), scan_archives=True)
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert "mongodb://user:pass123@" in hit.match_text

    def test_zip_txt_bearer_match_text(self, tmp_path: Path) -> None:
        """zip 内 txt 文件的 Bearer 令牌应被扫描到且 match_text 正确。"""
        zip_path = tmp_path / "tokens.zip"
        with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("auth.txt", f"Authorization: {_BEARER_SAMPLE}\n")

        scanner = Scanner(_build_ruleset(_bearer_rule()), scan_archives=True)
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "Bearer令牌")
        assert hit is not None
        assert hit.match_text == _BEARER_SAMPLE

    def test_zip_multiple_entries(self, tmp_path: Path) -> None:
        """zip 内多个文件应分别扫描，各自 match_text 独立。"""
        zip_path = tmp_path / "mixed.zip"
        with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("db.txt", f"url = {_DB_CONN_SAMPLE}\n")
            zf.writestr("auth.txt", f"Authorization: {_BEARER_SAMPLE}\n")
            zf.writestr("normal.txt", "nothing sensitive here\n")

        scanner = Scanner(
            _build_ruleset(_db_rule(), _bearer_rule()),
            scan_archives=True,
        )
        report = scanner.scan(tmp_path)
        db_hit = _find_hit(report.results, "数据库连接串")
        bearer_hit = _find_hit(report.results, "Bearer令牌")
        assert db_hit is not None
        assert bearer_hit is not None
        assert "mongodb://user:pass123@" in db_hit.match_text
        assert bearer_hit.match_text == _BEARER_SAMPLE

    def test_zip_nested_yaml_db_connection(self, tmp_path: Path) -> None:
        """zip 内 yaml 文件的数据库连接串应被扫描到。"""
        zip_path = tmp_path / "config.zip"
        with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("app.yaml", f"database:\n  url: {_DB_CONN_SAMPLE}\n")

        scanner = Scanner(_build_ruleset(_db_rule()), scan_archives=True)
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert "mongodb://user:pass123@" in hit.match_text

    def test_zip_db_connection_with_backslash(self, tmp_path: Path) -> None:
        """zip 内含反斜杠密码的数据库连接串 match_text 应原样保留。"""
        zip_path = tmp_path / "db.zip"
        with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("db.txt", f"url = {_DB_CONN_WITH_BACKSLASH}\n")

        scanner = Scanner(_build_ruleset(_db_rule()), scan_archives=True)
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        assert hit is not None
        assert "\\" in hit.match_text


# ---------------------------------------------------------------------------
# 二进制与不支持的格式
# ---------------------------------------------------------------------------


class TestBinaryAndUnsupportedFormats:
    """二进制文件与不支持格式的扫描安全性与 match_text 验证。"""

    def test_binary_file_no_crash_no_false_positive(self, tmp_path: Path) -> None:
        """二进制文件不应崩溃也不应误报。"""
        path = tmp_path / "data.bin"
        path.write_bytes(bytes(range(256)) * 4)
        scanner = Scanner(_build_ruleset(_db_rule(), _bearer_rule()))
        report = scanner.scan(tmp_path)
        assert len(report.hits) == 0

    def test_exe_file_no_crash(self, tmp_path: Path) -> None:
        """exe 文件不应崩溃。"""
        path = tmp_path / "app.exe"
        path.write_bytes(b"MZ\x90\x00" + b"\x00" * 512 + b"some binary content")
        scanner = Scanner(_build_ruleset(_db_rule(), _bearer_rule()))
        report = scanner.scan(tmp_path)
        assert len(report.hits) == 0

    def test_doc_unsupported_no_crash(self, tmp_path: Path) -> None:
        """旧版 .doc 格式（不支持）不应崩溃，无提取器时跳过内容。"""
        path = tmp_path / "old.doc"
        # 伪装的 OLE 文件头
        path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512)
        scanner = Scanner(_build_ruleset(_db_rule(), _bearer_rule()))
        report = scanner.scan(tmp_path)
        # 不支持提取的格式不会命中内容规则
        assert len(report.hits) == 0

    def test_xls_unsupported_no_crash(self, tmp_path: Path) -> None:
        """旧版 .xls 格式（不支持）不应崩溃。"""
        path = tmp_path / "old.xls"
        path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 512)
        scanner = Scanner(_build_ruleset(_db_rule(), _bearer_rule()))
        report = scanner.scan(tmp_path)
        assert len(report.hits) == 0

    def test_7z_unsupported_no_crash(self, tmp_path: Path) -> None:
        """7z 压缩格式（不支持）不应崩溃。"""
        path = tmp_path / "archive.7z"
        # 7z 魔数头
        path.write_bytes(b"7z\xbc\xaf\x27\x1c" + b"\x00" * 64)
        scanner = Scanner(
            _build_ruleset(_db_rule(), _bearer_rule()),
            scan_archives=True,
        )
        report = scanner.scan(tmp_path)
        # 无注册读取器的压缩格式会被跳过
        assert len(report.hits) == 0

    def test_tar_gz_unsupported_no_crash(self, tmp_path: Path) -> None:
        """tar.gz 压缩格式（不支持）不应崩溃。"""
        path = tmp_path / "archive.tar.gz"
        # gzip 魔数头 + 伪造内容
        path.write_bytes(b"\x1f\x8b\x08\x00" + b"\x00" * 64)
        scanner = Scanner(
            _build_ruleset(_db_rule(), _bearer_rule()),
            scan_archives=True,
        )
        report = scanner.scan(tmp_path)
        assert len(report.hits) == 0

    def test_binary_with_embedded_db_string(self, tmp_path: Path) -> None:
        """含二进制字节的 txt 文件中嵌入的数据库连接串应被 charset-normalizer 提取到。"""
        path = tmp_path / "mixed.txt"
        path.write_bytes(b"\x00\x01\x02" + _DB_CONN_SAMPLE.encode("utf-8") + b"\x03\x04\x05")
        scanner = Scanner(_build_ruleset(_db_rule()))
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "数据库连接串")
        # charset-normalizer 能从含二进制字节的文件中提取文本，应能命中
        assert hit is not None
        assert "mongodb://user:pass123@" in hit.match_text


# ---------------------------------------------------------------------------
# 内置规则集端到端验证
# ---------------------------------------------------------------------------


class TestBuiltinRulesetIntegration:
    """使用完整内置规则集的端到端扫描验证。"""

    def test_builtin_ruleset_scans_txt_password_and_key(self, tmp_path: Path) -> None:
        """内置规则集应同时扫描到 txt 中的密码赋值和私钥文件头。"""
        path = tmp_path / "secrets.txt"
        path.write_text(
            f"{_PASSWORD_SAMPLE}\n{_PRIVATE_KEY_SAMPLE}\n",
            encoding="utf-8",
        )
        ruleset = load_builtin_ruleset()
        scanner = Scanner(ruleset)
        report = scanner.scan(tmp_path)

        pwd_hit = _find_hit(report.results, "P0102-通用密码赋值")
        key_hit = _find_hit(report.results, "P0101-私钥文件头")
        assert pwd_hit is not None
        assert key_hit is not None
        assert pwd_hit.match_text == _PASSWORD_SAMPLE
        assert key_hit.match_text == _PRIVATE_KEY_SAMPLE

    def test_builtin_ruleset_scans_yaml_password(self, tmp_path: Path) -> None:
        """内置规则集应扫描到 yaml 中的密码赋值。"""
        path = tmp_path / "app.yaml"
        path.write_text(
            f"development:\n  database_{_PASSWORD_SAMPLE}\n",
            encoding="utf-8",
        )
        ruleset = load_builtin_ruleset()
        scanner = Scanner(ruleset)
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "P0102-通用密码赋值")
        assert hit is not None
        assert _PASSWORD_SAMPLE in hit.match_text

    def test_builtin_ruleset_scans_json_password(self, tmp_path: Path) -> None:
        """内置规则集应扫描到 json 中的密码赋值。"""
        path = tmp_path / "config.json"
        path.write_text(
            f'{{"config": "{_PASSWORD_SAMPLE}"}}',
            encoding="utf-8",
        )
        ruleset = load_builtin_ruleset()
        scanner = Scanner(ruleset)
        report = scanner.scan(tmp_path)
        hit = _find_hit(report.results, "P0102-通用密码赋值")
        assert hit is not None
        assert _PASSWORD_SAMPLE in hit.match_text
